"""File readers for different plan formats (.md, .txt, .docx, .pdf)"""

from pathlib import Path
from logger import logger


def read_plan_file(file_path: str) -> str:
    """
    Read a plan file and return its content as text.
    
    Supported formats:
    - .md, .txt: Plain text
    - .docx: Microsoft Word documents
    - .pdf: PDF documents
    
    Args:
        file_path: Path to the plan file
    
    Returns:
        Extracted text content
    
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    suffix = path.suffix.lower()
    
    if suffix in ('.md', '.txt'):
        return _read_text(path)
    elif suffix == '.docx':
        return _read_docx(path)
    elif suffix == '.pdf':
        return _read_pdf(path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Use .md, .txt, .docx, or .pdf")


def _read_text(path: Path) -> str:
    """Read plain text or markdown file."""
    logger.info(f"Reading text file: {path.name}")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _read_docx(path: Path) -> str:
    """Read Microsoft Word .docx file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    logger.info(f"Reading DOCX file: {path.name}")
    doc = Document(path)
    
    # Extract all paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Add markdown-style headings based on Word heading styles
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    
    return "\n\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    """Read PDF file."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    
    logger.info(f"Reading PDF file: {path.name}")
    reader = PdfReader(path)
    
    # Extract text from all pages
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
    
    if not pages:
        logger.warning("No text extracted from PDF")
        return ""
    
    return "\n\n".join(pages)
